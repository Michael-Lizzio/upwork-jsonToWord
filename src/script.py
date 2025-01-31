import os
import json
import yaml
from docx import Document


class JSONToWordConverter:
    def __init__(self, yaml_path, json_path, output_dir):
        self.yaml_path = yaml_path
        self.json_path = json_path
        self.output_dir = output_dir
        self.output_path = os.path.join(self.output_dir, "output.docx")
        self.static_text = self.load_yaml()
        self.data = self.load_json()
        self.doc = Document()

        # Ensure the output directory exists
        os.makedirs(self.output_dir, exist_ok=True)

    def load_yaml(self):
        """Loads static text from the YAML file."""
        if not os.path.exists(self.yaml_path):
            raise FileNotFoundError(f"YAML file not found: {self.yaml_path}")
        with open(self.yaml_path, "r") as yaml_file:
            return yaml.safe_load(yaml_file)

    def load_json(self):
        """Loads structured data from the JSON file."""
        if not os.path.exists(self.json_path):
            raise FileNotFoundError(f"JSON file not found: {self.json_path}")
        with open(self.json_path, "r") as json_file:
            return json.load(json_file)

    def create_table(self, title, headers, rows):
        """Creates a formatted table in the Word document."""
        self.doc.add_paragraph(title, style="Heading 2")
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = "Table Grid"

        # Add headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        # Add row data
        for row_idx, row_data in enumerate(rows, start=1):
            row_cells = table.rows[row_idx].cells
            for col_idx, value in enumerate(row_data):
                row_cells[col_idx].text = str(value)

        self.doc.add_paragraph("\n")

    def generate_document(self):
        """Generates the Word document based on JSON data."""
        segment_definitions = {
            "segment1": ["name", "value", "description"],
            "segment2": ["name", "value", "description"],
            "segment3": ["ID", "Score", "Category"],
            "segment4": ["Product", "Price", "Stock"],
            "segment5": ["Location", "Status", "Last Updated"]
        }

        for segment, headers in segment_definitions.items():
            if segment in self.data:
                self.doc.add_heading(self.static_text["segments"].get(segment, segment), level=1)

                for item in self.data[segment]:
                    for table_title in self.static_text["tables"].values():
                        rows = [[item.get(key, "N/A") for key in headers]]
                        self.create_table(table_title, headers, rows)

    def save_document(self):
        """Deletes old file if exists and saves the generated Word document."""
        if os.path.exists(self.output_path):
            try:
                os.remove(self.output_path)
                print("Previous output.docx deleted successfully.")
            except PermissionError:
                print("Error: output.docx is in use. Close the file and retry.")
                exit(1)

        self.doc.save(self.output_path)
        print(f"Word document saved successfully at: {self.output_path}")


def main():
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    DATA_DIR = os.path.abspath(os.path.join(BASE_DIR, "..", "data"))
    OUTPUT_DIR = os.path.abspath(os.path.join(BASE_DIR, "..", "output"))

    yaml_path = os.path.join(DATA_DIR, "static_text.yaml")
    json_path = os.path.join(DATA_DIR, "example.json")

    converter = JSONToWordConverter(yaml_path, json_path, OUTPUT_DIR)
    converter.generate_document()
    converter.save_document()


if __name__ == "__main__":
    main()
