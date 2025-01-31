import unittest
import json
from docx import Document

class TestDocxGeneration(unittest.TestCase):
    def setUp(self):
        with open("example.json", "r") as file:
            self.data = json.load(file)
        self.doc = Document("output.docx")

    def test_segment1_exists(self):
        """Check if segment 1 data exists in the document"""
        text = "\n".join([para.text for para in self.doc.paragraphs])
        for obj in self.data.get("segment1", []):
            self.assertIn(obj["name"], text)
            self.assertIn(str(obj["value"]), text)
            self.assertIn(obj["description"], text)

    def test_segment2_exists(self):
        """Check if segment 2 data exists in the document"""
        text = "\n".join([para.text for para in self.doc.paragraphs])
        for obj in self.data.get("segment2", []):
            self.assertIn(obj["name"], text)
            self.assertIn(str(obj["value"]), text)
            self.assertIn(obj["description"], text)

if __name__ == "__main__":
    unittest.main()
