import unittest
import os
import json
from docx import Document
from src.script import JSONToWordConverter


class TestDocxGeneration(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        """Runs once before all tests to generate the document."""
        BASE_DIR = os.path.dirname(os.path.abspath(__file__))
        DATA_DIR = os.path.abspath(os.path.join(BASE_DIR, "..", "data"))
        OUTPUT_DIR = os.path.abspath(os.path.join(BASE_DIR, "..", "output"))

        yaml_path = os.path.join(DATA_DIR, "static_text.yaml")
        json_path = os.path.join(DATA_DIR, "example.json")

        # Generate the document before running tests
        converter = JSONToWordConverter(yaml_path, json_path, OUTPUT_DIR)
        converter.generate_document()
        converter.save_document()

        with open(json_path, "r") as file:
            cls.data = json.load(file)
        cls.doc = Document(os.path.join(OUTPUT_DIR, "output.docx"))

    def extract_text_from_doc(self):
        """Extracts all text, including tables, from the document"""
        text = "\n".join([para.text for para in self.doc.paragraphs])

        # Extract table data
        for table in self.doc.tables:
            for row in table.rows:
                text += "\n" + "\t".join(cell.text.strip() for cell in row.cells)

        return text

    def test_segment1_exists(self):
        """Check if segment 1 data exists in the document"""
        text = self.extract_text_from_doc()
        for obj in self.data.get("segment1", []):
            self.assertIn(obj["name"], text)
            self.assertIn(str(obj["value"]), text)
            self.assertIn(obj["description"], text)

    def test_segment2_exists(self):
        """Check if segment 2 data exists in the document"""
        text = self.extract_text_from_doc()
        for obj in self.data.get("segment2", []):
            self.assertIn(obj["name"], text)
            self.assertIn(str(obj["value"]), text)
            self.assertIn(obj["description"], text)

    def test_document_contains_tables(self):
        """Ensure the generated document contains tables"""
        self.assertGreater(len(self.doc.tables), 0, "No tables found in the document")


if __name__ == "__main__":
    unittest.main()